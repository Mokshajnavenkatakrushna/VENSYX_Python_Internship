import os
import json
from google import genai
from pydantic import BaseModel, Field
from dotenv import load_dotenv

load_dotenv()
client = genai.Client()

GEMINI_SUPPORTED_MIME_TYPES = {
    ".pdf", ".jpg", ".jpeg", ".png", ".webp", ".gif",
    ".mp4", ".mov", ".avi", ".mp3", ".wav", ".txt"
}

# --- 1. Schemas for different bill types ---

class ElectricityBillSchema(BaseModel):
    flat_no: str = Field(description="House or flat number", default="")
    owner_name: str = Field(description="Owner or customer name", default="")
    account_no: str = Field(description="Account, RR, or consumer number", default="")
    bill_number: str = Field(description="Bill or invoice number", default="")
    bill_date: str = Field(description="Date of the bill", default="")
    due_date: str = Field(description="Due date of the bill", default="")
    period: str = Field(description="Billing period", default="")
    previous_reading: str = Field(description="Previous meter reading", default="")
    present_reading: str = Field(description="Present meter reading", default="")
    units: str = Field(description="Units consumed", default="")
    current_amount: str = Field(description="Current bill amount", default="")
    tax: str = Field(description="Tax amount", default="")
    fine: str = Field(description="Fine or late fee", default="")
    total_amount: str = Field(description="Total amount due", default="")
    company_name: str = Field(description="Electricity company name", default="")

class WaterBillSchema(BaseModel):
    flat_no: str = Field(description="House or flat number", default="")
    owner_name: str = Field(description="Owner or customer name", default="")
    water_connection_no: str = Field(description="Water connection or account number", default="")
    bill_number: str = Field(description="Bill or invoice number", default="")
    bill_date: str = Field(description="Date of the bill", default="")
    due_date: str = Field(description="Due date of the bill", default="")
    period: str = Field(description="Billing period", default="")
    consumption_litres: str = Field(description="Total consumption in litres or kilolitres", default="")
    water_charges: str = Field(description="Water usage charges", default="")
    sanitary_charges: str = Field(description="Sanitary or sewerage charges", default="")
    total_amount: str = Field(description="Total amount due", default="")
    company_name: str = Field(description="Water board or company name", default="")

class MaintenanceBillSchema(BaseModel):
    flat_no: str = Field(description="House or flat number", default="")
    owner_name: str = Field(description="Resident or owner name", default="")
    invoice_number: str = Field(description="Invoice number", default="")
    invoice_date: str = Field(description="Date of the invoice", default="")
    due_date: str = Field(description="Due date for payment", default="")
    maintenance_period: str = Field(description="Maintenance month or period", default="")
    maintenance_charges: str = Field(description="Base maintenance charges", default="")
    sinking_fund: str = Field(description="Sinking fund charges if any", default="")
    late_payment_penalty: str = Field(description="Penalty for late payment", default="")
    total_amount: str = Field(description="Total maintenance amount payable", default="")
    society_name: str = Field(description="Name of the apartment or housing society", default="")

class GenericBillSchema(BaseModel):
    invoice_number: str = Field(description="Invoice or bill number", default="")
    customer_name: str = Field(description="Customer name", default="")
    date: str = Field(description="Date of bill", default="")
    total_amount: str = Field(description="Total amount", default="")
    vendor_name: str = Field(description="Vendor or company name", default="")
    description: str = Field(description="Brief explanation of the bill", default="")

# Routing schema
class BillTypeCheck(BaseModel):
    bill_type: str = Field(description="Classify the document as: 'ELECTRICITY', 'WATER', 'MAINTENANCE', or 'OTHER'")

# --- 2. Processing logic ---

def _extract_docx_text(file_path: str) -> str:
    import docx
    doc = docx.Document(file_path)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    return "\n".join(lines)

def process_document_with_ai(file_path):
    if not os.path.exists(file_path):
        return "Error: File does not exist."

    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        # 1. First, classify the document type
        if ext in {".docx", ".doc"}:
            raw_text = _extract_docx_text(file_path)
            if not raw_text.strip():
                return "Error: Word document is empty."
            bill_type = _detect_type_from_text(raw_text)
            return _extract_data(bill_type, text_content=raw_text)
        else:
            uploaded_file = client.files.upload(file=file_path)
            bill_type = _detect_type_from_file(uploaded_file)
            result = _extract_data(bill_type, file_content=uploaded_file)
            client.files.delete(name=uploaded_file.name)
            return result

    except Exception as e:
        print(f"Error calling Gemini API: {e}")
        return f"Error During Extraction: {e}"

def _detect_type_from_file(uploaded_file):
    prompt = "Analyze this bill. Classify what type of bill this is (ELECTRICITY, WATER, MAINTENANCE, or OTHER)."
    response = client.models.generate_content(
        model='gemini-2.5-flash',
        contents=[uploaded_file, prompt],
        config={'response_mime_type': 'application/json', 'response_schema': BillTypeCheck, 'temperature': 0.1}
    )
    return json.loads(response.text).get('bill_type', 'OTHER')

def _detect_type_from_text(raw_text):
    prompt = f"Analyze this bill text. Classify what type of bill this is (ELECTRICITY, WATER, MAINTENANCE, or OTHER).\n\n{raw_text}"
    response = client.models.generate_content(
        model='gemini-2.5-flash',
        contents=[prompt],
        config={'response_mime_type': 'application/json', 'response_schema': BillTypeCheck, 'temperature': 0.1}
    )
    return json.loads(response.text).get('bill_type', 'OTHER')

def _extract_data(bill_type, file_content=None, text_content=None):
    # Select the schema based on our classification
    bill_type = str(bill_type).upper()
    if bill_type == 'ELECTRICITY':
        schema = ElectricityBillSchema
        title = "⚡ Electricity Bill"
    elif bill_type == 'WATER':
        schema = WaterBillSchema
        title = "💧 Water Bill"
    elif bill_type == 'MAINTENANCE':
        schema = MaintenanceBillSchema
        title = "🏢 Maintenance Bill"
    else:
        schema = GenericBillSchema
        title = "📄 Generic Bill"
        
    prompt = f"You are a strict data extractor. This is categorized as a {bill_type} bill. Extract the specific details into the exact schema provided. If a value is missing or not applicable, return an empty string."
    
    contents = [prompt]
    if file_content:
        contents.insert(0, file_content)
    else:
        contents.append(f"\nDocument Text:\n{text_content}")
        
    response = client.models.generate_content(
        model='gemini-2.5-flash',
        contents=contents,
        config={
            'response_mime_type': 'application/json',
            'response_schema': schema,
            'temperature': 0.1
        }
    )
    
    try:
        data = json.loads(response.text)
        
        # Requirement: print extracted JSON exactly as specified in the backend
        print(f"\n--- EXTRACTED {bill_type} JSON ---")
        print(json.dumps(data, indent=4))
        print("----------------------\n")
        
        # Requirement: simple output display type for the frontend (clean markdown list)
        markdown = f"### {title}\n\n"
        for k, v in data.items():
            formatted_key = k.replace('_', ' ').title()
            markdown += f"- **{formatted_key}**: {v}\n"
            
        return markdown
    except Exception as e:
        print(f"Failed to parse AI JSON: {e}")
        return f"```json\n{response.text}\n```"
